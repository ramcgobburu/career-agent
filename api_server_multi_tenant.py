"""
Multi-tenant FastAPI Backend Server for Career Agent
Supports user provisioning and context uploads for OpenAI Marketplace
"""

# Force immediate output for debugging
import sys
sys.stdout.flush()
print("🚀 MODULE LOADING: api_server_multi_tenant.py", flush=True)

import os
import tempfile
import logging
from pathlib import Path
from typing import Optional, List, Dict, Any
from datetime import datetime, timedelta
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Depends, Header, UploadFile, File, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, Field
from sqlalchemy.orm import Session

# Try to import Stripe
try:
    import stripe
    STRIPE_AVAILABLE = True
except ImportError:
    STRIPE_AVAILABLE = False
    stripe = None

from career_agent import CareerAgent
from database import (
    get_db, init_db, get_user_by_api_key, get_user_by_email, get_active_context_for_user,
    create_user, save_user_context, User, update_subscription, get_usage_limit_for_tier
)

logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Initialize database
init_db()

# Initialize FastAPI app
print("🚀 Creating FastAPI app...", flush=True)
app = FastAPI(
    title="Career Agent API - Multi-Tenant",
    description="AI-powered career materials generator API with user provisioning",
    version="2.0.0",
    servers=[
        {
            "url": os.getenv("API_BASE_URL", "https://career-agent-tf85.onrender.com"),
            "description": "Production server"
        }
    ]
)

# Enable CORS for frontend access
cors_origins = os.getenv("CORS_ORIGINS", "*")
if cors_origins == "*":
    allowed_origins = ["*"]
else:
    allowed_origins = [origin.strip() for origin in cors_origins.split(",")]

# Always allow the frontend domains
frontend_domains = [
    "https://careerpilotconsulting.com",
    "https://www.careerpilotconsulting.com",
    "http://localhost:3000",
    "http://localhost:3001"
]
for domain in frontend_domains:
    if domain not in allowed_origins:
        allowed_origins.append(domain)

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()

# Agent cache (per-user agents)
agent_cache: Dict[str, CareerAgent] = {}

# Stripe configuration
if STRIPE_AVAILABLE:
    stripe.api_key = os.getenv("STRIPE_SECRET_KEY")
    STRIPE_WEBHOOK_SECRET = os.getenv("STRIPE_WEBHOOK_SECRET")
    if not stripe.api_key:
        logger.warning("STRIPE_SECRET_KEY not set. Subscription features will be disabled.")
else:
    logger.warning("Stripe not installed. Install with: pip install stripe")
    STRIPE_WEBHOOK_SECRET = None


# Dependency to get current user
async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(security),
    db: Session = Depends(get_db)
) -> User:
    """Get current user from API key in Authorization header"""
    api_key = credentials.credentials
    user = get_user_by_api_key(db, api_key)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid API key")
    return user


# Alternative dependency for API key in header
async def get_current_user_from_header(
    x_api_key: Optional[str] = Header(None, alias="X-API-Key"),
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db)
) -> User:
    """Get current user from X-API-Key header or Authorization Bearer token (Supabase JWT)"""
    api_key = None
    token = None
    
    # Check X-API-Key header first
    if x_api_key:
        api_key = x_api_key
    # Check Authorization Bearer token
    elif authorization:
        if authorization.startswith("Bearer "):
            token = authorization.replace("Bearer ", "").strip()
        else:
            # If it's not Bearer, treat the whole value as API key
            api_key = authorization.strip()
    
    if not api_key and not token:
        raise HTTPException(
            status_code=401, 
            detail="API key required. Provide X-API-Key header or Authorization Bearer token"
        )
    
    # Try API key first (existing flow)
    if api_key:
        user = get_user_by_api_key(db, api_key)
        if user:
            return user
    
    # If token provided, try to get user from Supabase token
    if token:
        # Try to verify Supabase token and get user
        try:
            # Import Supabase if available
            from supabase import create_client, Client
            supabase_url = os.getenv("SUPABASE_URL")
            supabase_key = os.getenv("SUPABASE_ANON_KEY") or os.getenv("SUPABASE_SERVICE_ROLE_KEY")
            
            if supabase_url and supabase_key:
                supabase: Client = create_client(supabase_url, supabase_key)
                # Verify token and get user
                user_response = supabase.auth.get_user(token)
                if user_response and user_response.user:
                    supabase_user = user_response.user
                    email = supabase_user.email
                    
                    # Get or create user in our database
                    user = get_user_by_email(db, email)
                    if not user:
                        # Create user from Supabase info
                        from database import create_user
                        user = create_user(
                            db=db,
                            email=email,
                            name=supabase_user.user_metadata.get("name") or supabase_user.email.split("@")[0],
                            user_id=supabase_user.id
                        )
                    return user
        except ImportError:
            # Supabase not installed, fall through to try token as API key
            pass
        except Exception as e:
            logger.warning(f"Supabase token verification failed: {e}")
            # Fall through to try token as API key
            pass
        
        # If Supabase verification failed, try token as API key (backward compatibility)
        user = get_user_by_api_key(db, token)
        if user:
            return user
    
    raise HTTPException(status_code=401, detail="Invalid API key or authentication token")


def get_or_create_agent(user: User, db: Session) -> CareerAgent:
    """Get or create agent for user. Creates agent if context exists."""
    if user.id in agent_cache:
        return agent_cache[user.id]
    
    # Get user's context
    context = get_active_context_for_user(db, user.id)
    if not context:
        raise HTTPException(
            status_code=400,
            detail="No career context found. Please upload your career context first using /api/v1/upload-context"
        )
    
    # Create agent with user's context
    agent = CareerAgent(
        career_context_text=context.context_text,
        user_id=user.id,
        user_name=user.name or user.email or "the user",
        model_name=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        temperature=0.7
    )
    
    agent_cache[user.id] = agent
    return agent


# Request/Response Models
class CreateUserRequest(BaseModel):
    email: Optional[str] = Field(None, description="User email")
    name: Optional[str] = Field(None, description="User name")

class CreateUserResponse(BaseModel):
    user_id: str
    api_key: str
    message: str

class UploadContextRequest(BaseModel):
    context_text: Optional[str] = Field(None, description="Career context as text (alternative to file upload)")

class UploadContextResponse(BaseModel):
    success: bool
    message: str
    context_id: str

class CoverLetterRequest(BaseModel):
    company_name: str = Field(..., min_length=1, description="Name of the company")
    role_title: str = Field(..., min_length=1, description="Job title/role")
    job_description: Optional[str] = Field(None, description="Job description or requirements")
    additional_context: Optional[str] = Field(None, description="Additional context to consider")
    tone: str = Field("professional", description="Tone: professional, friendly, or formal")
    length: str = Field("medium", description="Length: short, medium, or long")
    format: Optional[str] = Field("text", description="Output format: text, markdown, or json")
    
class BlurbRequest(BaseModel):
    purpose: str = Field(..., description="Purpose of the blurb (e.g., 'LinkedIn introduction')")
    target_role: Optional[str] = Field(None, description="Target role to emphasize")
    max_words: int = Field(200, description="Maximum word count")
    style: str = Field("linkedin", description="Style: linkedin, email, or professional")
    format: Optional[str] = Field("text", description="Output format: text, markdown, or json")

class RoleSummaryRequest(BaseModel):
    role_type: str = Field(..., description="Role type (e.g., 'Product Manager')")
    focus_areas: Optional[List[str]] = Field(None, description="Areas to emphasize")
    format: Optional[str] = Field("text", description="Output format: text, markdown, or json")

class STARStoryRequest(BaseModel):
    project_name: Optional[str] = Field(None, description="Specific project name")
    situation_description: Optional[str] = Field(None, description="Situation to create story for")
    format: Optional[str] = Field("text", description="Output format: text, markdown, or json")

class InterviewAnswerRequest(BaseModel):
    question: str = Field(..., description="Interview question")
    company_context: Optional[str] = Field(None, description="Company-specific context")
    format: Optional[str] = Field("text", description="Output format: text, markdown, or json")

class QueryRequest(BaseModel):
    question: str = Field(..., description="Your question about career, skills, or experiences")
    format: Optional[str] = Field("text", description="Output format: text, markdown, or json")

class ResponseModel(BaseModel):
    success: bool
    content: str
    sources: Optional[List[Dict[str, Any]]] = None
    metadata: Optional[Dict[str, Any]] = None

class UserProfileResponse(BaseModel):
    id: str
    email: Optional[str]
    name: Optional[str]
    subscription_tier: str
    subscription_status: str
    requests_used: int
    requests_limit: int
    created_at: str
    is_active: bool

class UpdateUserRequest(BaseModel):
    name: Optional[str] = Field(None, description="User's full name")
    email: Optional[str] = Field(None, description="User's email address")

class DashboardStatsResponse(BaseModel):
    active_contexts: int = Field(0, description="Number of active/saved contexts")
    generated_documents: int = Field(0, description="Number of generated documents")
    job_applications: int = Field(0, description="Number of job applications")
    recent_activity: List[Dict[str, Any]] = Field(default_factory=list, description="Recent activity records")


def format_response(content: str, sources: Any, format_type: str = "text") -> str:
    """Format the response according to requested format."""
    if format_type == "json":
        import json
        return json.dumps({"content": content, "sources": str(sources)}, indent=2)
    elif format_type == "markdown":
        md = f"# Generated Content\n\n{content}\n\n"
        if sources:
            md += "## Sources\n\n"
            for i, source in enumerate(sources[:3], 1):
                md += f"### Source {i}\n{str(source)[:200]}...\n\n"
        return md
    else:  # text (default)
        return content


def convert_sources_to_dict(sources: Any) -> List[Dict[str, Any]]:
    """Convert LangChain Document objects to dictionaries."""
    if not sources:
        return []
    
    converted = []
    for source in sources:
        if hasattr(source, 'page_content') and hasattr(source, 'metadata'):
            converted.append({
                "page_content": source.page_content[:500],
                "metadata": source.metadata if source.metadata else {}
            })
        elif isinstance(source, dict):
            converted.append(source)
        else:
            converted.append({
                "content": str(source)[:500],
                "metadata": {}
            })
    return converted


# User Management Endpoints
@app.post("/api/v1/users", response_model=CreateUserResponse)
async def create_user_endpoint(
    request: CreateUserRequest,
    db: Session = Depends(get_db)
):
    """Create a new user and get an API key for authentication."""
    user = create_user(db, email=request.email, name=request.name)
    return CreateUserResponse(
        user_id=user.id,
        api_key=user.api_key,
        message="User created successfully. Save your API key - it won't be shown again!"
    )


# Job URL Parser Models
class ParseJobUrlRequest(BaseModel):
    url: str = Field(..., description="URL of the job posting (LinkedIn, Indeed, etc.)")

class ParseJobUrlResponse(BaseModel):
    success: bool
    company_name: Optional[str] = None
    role_title: Optional[str] = None
    job_description: Optional[str] = None
    error: Optional[str] = None

@app.post("/api/v1/parse-job-url", response_model=ParseJobUrlResponse)
async def parse_job_url(
    request: ParseJobUrlRequest,
    user: User = Depends(get_current_user_from_header)
):
    """Parse job posting URL and extract company name, role title, and job description using LLM."""
    try:
        import requests
        from bs4 import BeautifulSoup
        import re
        import json
        import os
        from openai import OpenAI
        
        # Validate URL
        if not request.url.startswith(('http://', 'https://')):
            return ParseJobUrlResponse(
                success=False,
                error="Invalid URL format. Please provide a full URL starting with http:// or https://"
            )
        
        # Set headers to mimic a browser
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
        }
        
        # Fetch the page
        response = requests.get(request.url, headers=headers, timeout=10)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove scripts, styles, and other non-content elements
        for script in soup(["script", "style", "nav", "header", "footer", "aside", "noscript"]):
            script.decompose()
        
        # Get the main content - try to find the most relevant content area
        main_content = None
        
        # Try to find main content areas
        content_selectors = [
            'main',
            'article',
            '[role="main"]',
            '.job-description',
            '.description',
            '#job-description',
            '.jobs-description',
            '.jobsearch-jobDescriptionText',
            'div.show-more-less-html__markup'
        ]
        
        for selector in content_selectors:
            elem = soup.select_one(selector)
            if elem:
                main_content = elem
                break
        
        # If no main content found, use body
        if not main_content:
            main_content = soup.find('body') or soup
        
        # Extract text content (limit to avoid token limits)
        page_text = main_content.get_text(separator='\n', strip=True)
        # Clean up excessive whitespace
        page_text = re.sub(r'\n{3,}', '\n\n', page_text)
        # Limit to first 8000 characters to stay within token limits
        page_text = page_text[:8000]
        
        # Also get page title for context
        page_title = soup.find('title')
        title_text = page_title.get_text(strip=True) if page_title else ""
        
        # Use OpenAI to intelligently extract information
        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            return ParseJobUrlResponse(
                success=False,
                error="OpenAI API key not configured. Cannot parse job posting."
            )
        
        client = OpenAI(api_key=openai_api_key)
        
        # Create a prompt for the LLM
        prompt = f"""You are analyzing a job posting webpage. Extract the following information from the content below:

Page Title: {title_text}

Page Content:
{page_text}

Please extract and return ONLY a JSON object with these exact keys:
- "company_name": The name of the hiring company (string, or null if not found)
- "role_title": The job title/position name (string, or null if not found)  
- "job_description": A clean, well-formatted job description including key responsibilities and requirements (string, or null if not found)

Rules:
1. Extract the actual company name, not the job board name (e.g., "Google" not "LinkedIn" or "Indeed")
2. Extract the full job title as it appears
3. For job_description, include the main responsibilities, requirements, and key details. Format it nicely with line breaks.
4. If any field cannot be determined, use null
5. Return ONLY valid JSON, no markdown, no explanations

JSON:"""
        
        try:
            # Call OpenAI API
            completion = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are a helpful assistant that extracts structured information from job postings. Always return valid JSON only."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                temperature=0.1,  # Low temperature for accurate extraction
                response_format={"type": "json_object"}  # Force JSON response
            )
            
            # Parse the JSON response
            response_text = completion.choices[0].message.content
            extracted_data = json.loads(response_text)
            
            company_name = extracted_data.get("company_name")
            role_title = extracted_data.get("role_title")
            job_description = extracted_data.get("job_description")
            
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM JSON response: {e}")
            return ParseJobUrlResponse(
                success=False,
                error="Failed to parse job posting. Please try again or enter details manually."
            )
        except Exception as e:
            logger.error(f"OpenAI API error: {e}")
            return ParseJobUrlResponse(
                success=False,
                error=f"Error analyzing job posting: {str(e)}"
            )
        
        # Clean up results
        if role_title:
            role_title = role_title.strip()
        if company_name:
            company_name = company_name.strip()
        if job_description:
            job_description = job_description.strip()
            # Limit description length
            if len(job_description) > 5000:
                job_description = job_description[:5000] + "..."
        
        # Check if we got at least some information
        if not role_title and not company_name and not job_description:
            return ParseJobUrlResponse(
                success=False,
                error="Could not extract job information from this URL. Please try a different job posting or enter details manually."
            )
        
        return ParseJobUrlResponse(
            success=True,
            company_name=company_name,
            role_title=role_title,
            job_description=job_description
        )
        
    except requests.exceptions.RequestException as e:
        return ParseJobUrlResponse(
            success=False,
            error=f"Failed to fetch job posting: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error parsing job URL: {str(e)}")
        return ParseJobUrlResponse(
            success=False,
            error=f"Error parsing job posting: {str(e)}"
        )


@app.post("/api/v1/upload-context", response_model=UploadContextResponse)
async def upload_context(
    file: Optional[UploadFile] = File(None),
    context_text: Optional[str] = Form(None),
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Upload career context document (text, markdown, doc, or pdf file)."""
    # Clear agent cache for this user
    if user.id in agent_cache:
        del agent_cache[user.id]
    
    # Get context text from file or form
    if file:
        file_name = file.filename
        file_type = Path(file_name).suffix.lower() if file_name else None
        content = await file.read()
        
        # Handle different file types
        if file_type in ['.txt', '.md', '.markdown']:
            # Text files
            if file.content_type not in ["text/plain", "text/markdown", "text/x-markdown", None]:
                raise HTTPException(status_code=400, detail="File must be text or markdown")
            context_text = content.decode('utf-8')
        elif file_type in ['.pdf']:
            # PDF files
            try:
                import pypdf
                from io import BytesIO
                pdf_reader = pypdf.PdfReader(BytesIO(content))
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                context_text = '\n\n'.join(text_parts)
            except ImportError:
                raise HTTPException(
                    status_code=503,
                    detail="PDF support not available. Please install pypdf: pip install pypdf"
                )
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
        elif file_type in ['.doc', '.docx']:
            # Word documents
            try:
                from docx import Document as DocxDocument
                from io import BytesIO
                doc = DocxDocument(BytesIO(content))
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                context_text = '\n\n'.join(text_parts)
            except ImportError:
                raise HTTPException(
                    status_code=503,
                    detail="Word document support not available. Please install python-docx: pip install python-docx"
                )
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading Word document: {str(e)}")
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_type}. Supported types: .txt, .md, .pdf, .doc, .docx"
            )
    elif context_text:
        file_name = None
        file_type = None
    else:
        raise HTTPException(status_code=400, detail="Either file or context_text must be provided")
    
    if not context_text or len(context_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Context text is too short (minimum 50 characters)")
    
    # Save context
    context = save_user_context(
        db,
        user_id=user.id,
        context_text=context_text,
        file_name=file_name,
        file_type=file_type
    )
    
    return UploadContextResponse(
        success=True,
        message="Career context uploaded successfully",
        context_id=context.id
    )


# API Information
@app.get("/")
@app.post("/api/v1/upload-context", response_model=UploadContextResponse)
async def upload_context(
    file: Optional[UploadFile] = File(None),
    context_text: Optional[str] = Form(None),
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Upload career context document (text, markdown, doc, or pdf file)."""
    # Clear agent cache for this user
    if user.id in agent_cache:
        del agent_cache[user.id]
    
    # Get context text from file or form
    if file:
        file_name = file.filename
        file_type = Path(file_name).suffix.lower() if file_name else None
        content = await file.read()
        
        # Handle different file types
        if file_type in ['.txt', '.md', '.markdown']:
            # Text files
            if file.content_type not in ["text/plain", "text/markdown", "text/x-markdown", None]:
                raise HTTPException(status_code=400, detail="File must be text or markdown")
            context_text = content.decode('utf-8')
        elif file_type in ['.pdf']:
            # PDF files
            try:
                import pypdf
                from io import BytesIO
                pdf_reader = pypdf.PdfReader(BytesIO(content))
                text_parts = []
                for page in pdf_reader.pages:
                    text_parts.append(page.extract_text())
                context_text = '\n\n'.join(text_parts)
            except ImportError:
                raise HTTPException(
                    status_code=503,
                    detail="PDF support not available. Please install pypdf: pip install pypdf"
                )
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
        elif file_type in ['.doc', '.docx']:
            # Word documents
            try:
                from docx import Document as DocxDocument
                from io import BytesIO
                doc = DocxDocument(BytesIO(content))
                text_parts = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                context_text = '\n\n'.join(text_parts)
            except ImportError:
                raise HTTPException(
                    status_code=503,
                    detail="Word document support not available. Please install python-docx: pip install python-docx"
                )
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading Word document: {str(e)}")
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_type}. Supported types: .txt, .md, .pdf, .doc, .docx"
            )
    elif context_text:
        file_name = None
        file_type = None
    else:
        raise HTTPException(status_code=400, detail="Either file or context_text must be provided")
    
    if not context_text or len(context_text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Context text is too short (minimum 50 characters)")
    
    # Save context
    context = save_user_context(
        db,
        user_id=user.id,
        context_text=context_text,
        file_name=file_name,
        file_type=file_type
    )
    
    return UploadContextResponse(
        success=True,
        message="Career context uploaded successfully",
        context_id=context.id
    )


# API Information
@app.get("/")
async def root():
    """Root endpoint - API information."""
    return {
        "message": "Career Agent API - Multi-Tenant",
        "version": "2.0.0",
        "status": "running",
        "endpoints": {
            "create_user": "/api/v1/users",
            "upload_context": "/api/v1/upload-context",
            "cover_letter": "/api/v1/cover-letter",
            "blurb": "/api/v1/blurb",
            "role_summary": "/api/v1/role-summary",
            "star_story": "/api/v1/star-story",
            "interview_answer": "/api/v1/interview-answer",
            "query": "/api/v1/query"
        },
        "authentication": "Use X-API-Key header or Authorization: Bearer <api_key>",
        "docs": "/docs"
    }


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {
        "status": "healthy",
        "database": "connected",
        "agent_cache_size": len(agent_cache)
    }


@app.get("/privacy-policy")
async def privacy_policy():
    """Serve privacy policy page for OpenAI Actions requirement."""
    return FileResponse("privacy_policy.html", media_type="text/html")


# Career Generation Endpoints (require authentication)
@app.post("/api/v1/cover-letter", response_model=ResponseModel)
async def generate_cover_letter(
    request: CoverLetterRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Generate a personalized cover letter."""
    logger.info(f"Cover letter request from user {user.id}: company={request.company_name}, role={request.role_title}")
    agent = get_or_create_agent(user, db)
    
    try:
        result = agent.generate_cover_letter(
            company_name=request.company_name,
            role_title=request.role_title,
            job_description=request.job_description,
            additional_context=request.additional_context,
            tone=request.tone,
            length=request.length
        )
        
        formatted_content = format_response(
            result["content"],
            result.get("sources", []),
            request.format or "text"
        )
        
        converted_sources = convert_sources_to_dict(result.get("sources", []))
        
        return ResponseModel(
            success=True,
            content=formatted_content,
            sources=converted_sources,
            metadata={
                "company": request.company_name,
                "role": request.role_title,
                "tone": request.tone,
                "length": request.length,
                "has_additional_context": bool(request.additional_context)
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating cover letter: {str(e)}")


@app.post("/api/v1/blurb", response_model=ResponseModel)
async def generate_blurb(
    request: BlurbRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Generate a short blurb for LinkedIn, email, etc."""
    agent = get_or_create_agent(user, db)
    
    try:
        result = agent.generate_blurb(
            purpose=request.purpose,
            target_role=request.target_role,
            max_words=request.max_words,
            style=request.style
        )
        
        formatted_content = format_response(
            result["content"],
            result.get("sources", []),
            request.format or "text"
        )
        
        converted_sources = convert_sources_to_dict(result.get("sources", []))
        
        return ResponseModel(
            success=True,
            content=formatted_content,
            sources=converted_sources,
            metadata={
                "purpose": request.purpose,
                "target_role": request.target_role,
                "max_words": request.max_words,
                "style": request.style
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating blurb: {str(e)}")


@app.post("/api/v1/role-summary", response_model=ResponseModel)
async def generate_role_summary(
    request: RoleSummaryRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Generate a role-specific professional summary."""
    agent = get_or_create_agent(user, db)
    
    try:
        result = agent.generate_role_specific_summary(
            role_type=request.role_type,
            focus_areas=request.focus_areas
        )
        
        formatted_content = format_response(
            result["content"],
            result.get("sources", []),
            request.format or "text"
        )
        
        converted_sources = convert_sources_to_dict(result.get("sources", []))
        
        return ResponseModel(
            success=True,
            content=formatted_content,
            sources=converted_sources,
            metadata={
                "role_type": request.role_type,
                "focus_areas": request.focus_areas
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating role summary: {str(e)}")


@app.post("/api/v1/star-story", response_model=ResponseModel)
async def generate_star_story(
    request: STARStoryRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Generate or retrieve a STAR story."""
    agent = get_or_create_agent(user, db)
    
    try:
        result = agent.generate_star_story(
            project_name=request.project_name,
            situation_description=request.situation_description
        )
        
        formatted_content = format_response(
            result["content"],
            result.get("sources", []),
            request.format or "text"
        )
        
        converted_sources = convert_sources_to_dict(result.get("sources", []))
        
        return ResponseModel(
            success=True,
            content=formatted_content,
            sources=converted_sources,
            metadata={
                "project_name": request.project_name,
                "situation_description": request.situation_description
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating STAR story: {str(e)}")


@app.post("/api/v1/interview-answer", response_model=ResponseModel)
async def answer_interview_question(
    request: InterviewAnswerRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Generate an answer to an interview question."""
    agent = get_or_create_agent(user, db)
    
    try:
        result = agent.answer_interview_question(
            question=request.question,
            company_context=request.company_context
        )
        
        formatted_content = format_response(
            result["content"],
            result.get("sources", []),
            request.format or "text"
        )
        
        converted_sources = convert_sources_to_dict(result.get("sources", []))
        
        return ResponseModel(
            success=True,
            content=formatted_content,
            sources=converted_sources,
            metadata={
                "question": request.question,
                "company_context": request.company_context
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating interview answer: {str(e)}")


@app.post("/api/v1/query", response_model=ResponseModel)
async def query_agent(
    request: QueryRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Generic query endpoint for any career-related question."""
    agent = get_or_create_agent(user, db)
    
    try:
        result = agent.query(request.question)
        
        formatted_content = format_response(
            result["content"],
            result.get("sources", []),
            request.format or "text"
        )
        
        converted_sources = convert_sources_to_dict(result.get("sources", []))
        
        return ResponseModel(
            success=True,
            content=formatted_content,
            sources=converted_sources,
            metadata={
                "question": request.question
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing query: {str(e)}")


# Subscription Models
class CreateCheckoutSessionRequest(BaseModel):
    price_id: str = Field(..., description="Stripe Price ID for the subscription plan")
    success_url: Optional[str] = Field(None, description="URL to redirect after successful payment")
    cancel_url: Optional[str] = Field(None, description="URL to redirect after cancelled payment")

class CheckoutSessionResponse(BaseModel):
    session_id: str
    url: str

class SubscriptionStatusResponse(BaseModel):
    subscription_tier: str
    subscription_status: str
    subscription_expires_at: Optional[str]
    requests_used: int
    requests_limit: int
    stripe_customer_id: Optional[str] = None
    stripe_subscription_id: Optional[str] = None


# Subscription Endpoints
print("🔍 DEBUG: About to define subscription routes...")
@app.post("/api/v1/create-checkout-session", response_model=CheckoutSessionResponse)
async def create_checkout_session(
    request: CreateCheckoutSessionRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Create a Stripe Checkout Session for subscription."""
    if not STRIPE_AVAILABLE or not stripe.api_key:
        raise HTTPException(
            status_code=503,
            detail="Payment processing not configured. Please contact support."
        )
    
    try:
        # Get base URL from environment or request
        base_url = os.getenv("FRONTEND_URL", "https://careerpilotconsulting.com")
        success_url = request.success_url or f"{base_url}/subscription?success=true"
        cancel_url = request.cancel_url or f"{base_url}/subscription?canceled=true"
        
        # Create or retrieve Stripe customer
        customer_email = user.email
        if not customer_email:
            raise HTTPException(
                status_code=400,
                detail="User email is required for subscription. Please update your profile."
            )
        
        # Try to find existing customer by email
        customers = stripe.Customer.list(email=customer_email, limit=1)
        if customers.data:
            customer_id = customers.data[0].id
        else:
            # Create new customer
            customer = stripe.Customer.create(
                email=customer_email,
                name=user.name,
                metadata={"user_id": user.id}
            )
            customer_id = customer.id
        
        # Create checkout session
        checkout_session = stripe.checkout.Session.create(
            customer=customer_id,
            payment_method_types=["card"],
            line_items=[
                {
                    "price": request.price_id,
                    "quantity": 1,
                }
            ],
            mode="subscription",
            success_url=success_url,
            cancel_url=cancel_url,
            metadata={
                "user_id": user.id,
                "user_email": customer_email
            },
            allow_promotion_codes=True,
        )
        
        return CheckoutSessionResponse(
            session_id=checkout_session.id,
            url=checkout_session.url
        )
    
    except Exception as e:
        # Check if it's a Stripe error
        error_type = type(e).__name__
        error_str = str(type(e)).lower()
        if "Stripe" in error_type or "stripe" in error_str or "AuthenticationError" in error_type:
            logger.error(f"Stripe error creating checkout session: {e}")
            error_message = str(e)
            if "Invalid API Key" in error_message:
                raise HTTPException(
                    status_code=503,
                    detail="Payment processing not configured correctly. Please contact support."
                )
            raise HTTPException(
                status_code=400,
                detail=f"Payment processing error: {error_message}"
            )
        # Generic error handler
        logger.error(f"Error creating checkout session: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail=f"Error creating checkout session: {str(e)}"
        )


print("🔍 DEBUG: About to define subscription-status route...")
@app.get("/api/v1/subscription-status", response_model=SubscriptionStatusResponse)
async def get_subscription_status(
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Get current subscription status for the user."""
    limit = get_usage_limit_for_tier(user.subscription_tier)
    
    expires_at_str = None
    if user.subscription_expires_at:
        expires_at_str = user.subscription_expires_at.isoformat()
    
    # Try to get Stripe subscription info if available
    stripe_customer_id = None
    stripe_subscription_id = None
    
    if STRIPE_AVAILABLE and stripe.api_key and user.email:
        try:
            customers = stripe.Customer.list(email=user.email, limit=1)
            if customers.data:
                customer = customers.data[0]
                stripe_customer_id = customer.id
                
                # Get active subscriptions
                subscriptions = stripe.Subscription.list(
                    customer=customer.id,
                    status="active",
                    limit=1
                )
                if subscriptions.data:
                    stripe_subscription_id = subscriptions.data[0].id
        except Exception as e:
            logger.warning(f"Could not fetch Stripe subscription info: {e}")
    
    return SubscriptionStatusResponse(
        subscription_tier=user.subscription_tier,
        subscription_status=user.subscription_status,
        subscription_expires_at=expires_at_str,
        requests_used=user.requests_used,
        requests_limit=limit,
        stripe_customer_id=stripe_customer_id,
        stripe_subscription_id=stripe_subscription_id
    )


# Account Management Endpoints
@app.get("/api/v1/dashboard-stats", response_model=DashboardStatsResponse)
async def get_dashboard_stats(
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Get dashboard statistics for the current user."""
    from database import (
        list_user_contexts,
        list_generated_documents,
        get_recent_usage_records
    )
    
    # Get active contexts count
    contexts = list_user_contexts(db, user.id, limit=1000)
    active_contexts = len([c for c in contexts if c.is_active])
    
    # Get generated documents count
    documents = list_generated_documents(db, user.id, limit=1000)
    generated_documents = len(documents)
    
    # Count resumes (documents with type containing 'resume')
    resumes = [d for d in documents if 'resume' in d.document_type.lower()]
    active_resumes = len(resumes)
    
    # For now, job applications = generated documents with type 'job-application' or similar
    job_applications = len([d for d in documents if 'application' in d.document_type.lower() or 'cover-letter' in d.document_type.lower()])
    
    # Get recent activity (last 10 usage records)
    recent_records = get_recent_usage_records(db, user.id, limit=10)
    recent_activity = []
    
    for record in recent_records:
        # Map endpoint to friendly action name
        action_map = {
            'cover-letter': 'Generated cover letter',
            'blurb': 'Generated blurb',
            'role-summary': 'Generated role summary',
            'star-story': 'Generated STAR story',
            'interview-answer': 'Generated interview answer',
            'query': 'Ran query',
            'upload-context': 'Uploaded context'
        }
        action = action_map.get(record.endpoint, f'Used {record.endpoint}')
        
        # Calculate time ago
        time_ago = get_time_ago(record.created_at)
        
        recent_activity.append({
            'action': action,
            'endpoint': record.endpoint,
            'time_ago': time_ago,
            'created_at': record.created_at.isoformat() if record.created_at else None
        })
    
    return DashboardStatsResponse(
        active_contexts=len(contexts),  # Total saved contexts
        generated_documents=generated_documents,
        job_applications=job_applications,
        recent_activity=recent_activity
    )


def get_time_ago(dt: datetime) -> str:
    """Convert datetime to human-readable time ago string."""
    if not dt:
        return 'Unknown'
    
    now = datetime.utcnow()
    diff = now - dt
    
    if diff.days > 0:
        return f'{diff.days} day{"s" if diff.days > 1 else ""} ago'
    elif diff.seconds >= 3600:
        hours = diff.seconds // 3600
        return f'{hours} hour{"s" if hours > 1 else ""} ago'
    elif diff.seconds >= 60:
        minutes = diff.seconds // 60
        return f'{minutes} minute{"s" if minutes > 1 else ""} ago'
    else:
        return 'Just now'


@app.get("/api/v1/me", response_model=UserProfileResponse)
async def get_user_profile(
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Get current user's profile information."""
    from database import get_usage_limit_for_tier
    limit = get_usage_limit_for_tier(user.subscription_tier)
    
    return UserProfileResponse(
        id=user.id,
        email=user.email,
        name=user.name,
        subscription_tier=user.subscription_tier,
        subscription_status=user.subscription_status,
        requests_used=user.requests_used,
        requests_limit=limit,
        created_at=user.created_at.isoformat() if user.created_at else "",
        is_active=user.is_active
    )


@app.put("/api/v1/me", response_model=UserProfileResponse)
async def update_user_profile(
    request: UpdateUserRequest,
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Update user's profile (name and/or email)."""
    if request.name is not None:
        user.name = request.name.strip() if request.name else None
    if request.email is not None:
        # Check if email is already taken by another user
        existing_user = db.query(User).filter(
            User.email == request.email.strip(),
            User.id != user.id
        ).first()
        if existing_user:
            raise HTTPException(status_code=400, detail="Email already in use")
        user.email = request.email.strip() if request.email else None
    
    user.updated_at = datetime.now()
    db.commit()
    db.refresh(user)
    
    from database import get_usage_limit_for_tier
    limit = get_usage_limit_for_tier(user.subscription_tier)
    
    return UserProfileResponse(
        id=user.id,
        email=user.email,
        name=user.name,
        subscription_tier=user.subscription_tier,
        subscription_status=user.subscription_status,
        requests_used=user.requests_used,
        requests_limit=limit,
        created_at=user.created_at.isoformat() if user.created_at else "",
        is_active=user.is_active
    )


@app.post("/api/v1/me/pause")
async def pause_account(
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Pause user account (set is_active to False)."""
    user.is_active = False
    user.updated_at = datetime.now()
    db.commit()
    
    return {
        "success": True,
        "message": "Account paused successfully. You can reactivate by contacting support."
    }


@app.delete("/api/v1/me")
async def delete_account(
    user: User = Depends(get_current_user_from_header),
    db: Session = Depends(get_db)
):
    """Delete user account and all associated data."""
    # Delete all related data (cascade should handle this, but being explicit)
    user_id = user.id
    
    # Clear agent cache
    if user_id in agent_cache:
        del agent_cache[user_id]
    
    # Delete user (cascade will delete contexts, usage records, etc.)
    db.delete(user)
    db.commit()
    
    return {
        "success": True,
        "message": "Account deleted successfully. All your data has been removed."
    }


@app.post("/api/v1/webhooks/stripe")
async def stripe_webhook(request: Request):
    """Handle Stripe webhook events for subscription management."""
    if not STRIPE_AVAILABLE or not STRIPE_WEBHOOK_SECRET:
        raise HTTPException(status_code=503, detail="Webhook secret not configured")
    
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")
    
    if not sig_header:
        raise HTTPException(status_code=400, detail="Missing stripe-signature header")
    
    try:
        event = stripe.Webhook.construct_event(
            payload, sig_header, STRIPE_WEBHOOK_SECRET
        )
    except ValueError as e:
        logger.error(f"Invalid payload: {e}")
        raise HTTPException(status_code=400, detail="Invalid payload")
    except Exception as e:
        # Check if it's a Stripe signature verification error
        error_type = type(e).__name__
        if "SignatureVerificationError" in error_type or "stripe" in str(type(e)).lower():
            logger.error(f"Invalid signature: {e}")
            raise HTTPException(status_code=400, detail="Invalid signature")
        # Re-raise other exceptions
        raise
    
    # Handle webhook events
    event_type = event["type"]
    event_data = event["data"]["object"]
    
    logger.info(f"Received Stripe webhook: {event_type}")
    
    db = next(get_db())
    
    try:
        if event_type == "checkout.session.completed":
            # Payment succeeded - upgrade subscription
            session = event_data
            user_id = session.get("metadata", {}).get("user_id")
            customer_email = session.get("customer_email") or session.get("customer_details", {}).get("email")
            
            if user_id:
                # Find user by ID
                user = db.query(User).filter(User.id == user_id).first()
            elif customer_email:
                # Find user by email
                from database import get_user_by_email
                user = get_user_by_email(db, customer_email)
            else:
                logger.warning("No user_id or email in checkout session")
                return {"status": "ok"}
            
            if user:
                # Set subscription to premium with 30 days expiration
                expires_at = datetime.now() + timedelta(days=30)
                update_subscription(db, user.id, "premium", expires_at)
                logger.info(f"Upgraded user {user.id} to premium via webhook")
            else:
                logger.warning(f"User not found for checkout session: {user_id or customer_email}")
        
        elif event_type == "customer.subscription.created":
            # New subscription created
            subscription = event_data
            customer_id = subscription.get("customer")
            
            if customer_id:
                customer = stripe.Customer.retrieve(customer_id)
                customer_email = customer.email
                
                if customer_email:
                    from database import get_user_by_email
                    user = get_user_by_email(db, customer_email)
                    if user:
                        expires_at = datetime.fromtimestamp(subscription.get("current_period_end", 0))
                        update_subscription(db, user.id, "premium", expires_at)
                        logger.info(f"Subscription created for user {user.id}")
        
        elif event_type == "customer.subscription.updated":
            # Subscription updated (renewed, changed, etc.)
            subscription = event_data
            customer_id = subscription.get("customer")
            status = subscription.get("status")
            
            if customer_id:
                customer = stripe.Customer.retrieve(customer_id)
                customer_email = customer.email
                
                if customer_email:
                    from database import get_user_by_email
                    user = get_user_by_email(db, customer_email)
                    if user:
                        if status in ["active", "trialing"]:
                            expires_at = datetime.fromtimestamp(subscription.get("current_period_end", 0))
                            update_subscription(db, user.id, "premium", expires_at)
                            logger.info(f"Subscription updated for user {user.id}")
                        elif status in ["canceled", "unpaid", "past_due"]:
                            # Downgrade to free
                            update_subscription(db, user.id, "free", None)
                            logger.info(f"Subscription canceled for user {user.id}")
        
        elif event_type == "customer.subscription.deleted":
            # Subscription cancelled
            subscription = event_data
            customer_id = subscription.get("customer")
            
            if customer_id:
                customer = stripe.Customer.retrieve(customer_id)
                customer_email = customer.email
                
                if customer_email:
                    from database import get_user_by_email
                    user = get_user_by_email(db, customer_email)
                    if user:
                        update_subscription(db, user.id, "free", None)
                        logger.info(f"Subscription deleted for user {user.id}")
        
        elif event_type == "invoice.payment_failed":
            # Payment failed
            invoice = event_data
            customer_id = invoice.get("customer")
            
            if customer_id:
                customer = stripe.Customer.retrieve(customer_id)
                customer_email = customer.email
                
                if customer_email:
                    from database import get_user_by_email
                    user = get_user_by_email(db, customer_email)
                    if user:
                        logger.warning(f"Payment failed for user {user.id}")
                        # Optionally downgrade or mark subscription as expired
                        # For now, just log it
        
        return {"status": "success"}
    
    except Exception as e:
        logger.error(f"Error processing webhook: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Error processing webhook")
    finally:
        db.close()


# Log registered routes on startup - using print for visibility
@app.on_event("startup")
async def log_routes():
    """Log all registered routes for debugging."""
    routes = []
    for route in app.routes:
        if hasattr(route, 'path') and hasattr(route, 'methods'):
            method = list(route.methods)[0] if route.methods else 'GET'
            routes.append(f"{method} {route.path}")
    
    print(f"✅ Registered {len(routes)} routes")
    logger.info(f"✅ Registered {len(routes)} routes")
    subscription_routes = [r for r in routes if 'subscription' in r.lower() or 'checkout' in r.lower()]
    if subscription_routes:
        print(f"✅ Subscription routes: {', '.join(subscription_routes)}")
        logger.info(f"✅ Subscription routes: {', '.join(subscription_routes)}")
    else:
        print("⚠️  No subscription routes found!")
        logger.warning("⚠️  No subscription routes found!")
    
    # Also print all /api/v1 routes for debugging
    api_routes = [r for r in routes if '/api/v1' in r]
    print(f"📋 All /api/v1 routes ({len(api_routes)}):")
    for r in sorted(api_routes):
        print(f"   {r}")


if __name__ == "__main__":
    import uvicorn
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", 8000))
    reload = os.getenv("ENVIRONMENT", "development") == "development"
    
    uvicorn.run(
        "api_server_multi_tenant:app",
        host=host,
        port=port,
        reload=reload,
        log_level="info"
    )


